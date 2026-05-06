from pathlib import Path
from textwrap import wrap

from PIL import Image, ImageDraw, ImageFont, ImageFilter
from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(r"C:\Users\amare\controle-equipamento")
OUT_DIR = ROOT / "portfolio"
DOCX_PATH = OUT_DIR / "Portfolio_Equipamento_Gestao_v6.docx"
LOGO_PATH = ROOT / "app" / "src" / "assets" / "logo-sistema.png"
COVER_LOGO_PATH = OUT_DIR / "portfolio_cover_logo.png"
SCREENSHOT_DIR = OUT_DIR / "prints"
PROCESSED_DIR = OUT_DIR / "prints_processados"
IMAGE_WIDTH = Inches(7.1)

PRIMARY = "#0B5ED7"
PRIMARY_DARK = "#163256"
PRIMARY_SOFT = "#EAF2FF"
GREEN = "#19A463"
ORANGE = "#FF9F1C"
PURPLE = "#6F42C1"
GRAY = "#5F6F86"
RED = "#D6336C"


def hex_rgb(value: str):
    value = value.lstrip("#")
    return tuple(int(value[i : i + 2], 16) for i in (0, 2, 4))


def font(size: int, bold: bool = False):
    candidates = [
        ("arialbd.ttf" if bold else "arial.ttf"),
        ("C:/Windows/Fonts/arialbd.ttf" if bold else "C:/Windows/Fonts/arial.ttf"),
        ("C:/Windows/Fonts/calibrib.ttf" if bold else "C:/Windows/Fonts/calibri.ttf"),
    ]
    for candidate in candidates:
        try:
            return ImageFont.truetype(candidate, size=size)
        except OSError:
            continue
    return ImageFont.load_default()


def draw_wrapped(draw, text, box, fill, font_obj, line_gap=6, align="left"):
    x, y, w, _h = box
    avg = max(font_obj.size * 0.58, 1)
    width_chars = max(int(w / avg), 8)
    lines = []
    for paragraph in text.split("\n"):
        lines.extend(wrap(paragraph, width=width_chars) or [""])
    for line in lines:
        bbox = draw.textbbox((0, 0), line, font=font_obj)
        line_w = bbox[2] - bbox[0]
        line_h = bbox[3] - bbox[1]
        line_x = x
        if align == "center":
            line_x = x + (w - line_w) / 2
        elif align == "right":
            line_x = x + (w - line_w)
        draw.text((line_x, y), line, fill=fill, font=font_obj)
        y += line_h + line_gap
    return y


def rounded(draw, xy, radius, fill, outline=None, width=1):
    draw.rounded_rectangle(xy, radius=radius, fill=fill, outline=outline, width=width)


def create_cover_logo(path: Path):
    img = Image.open(LOGO_PATH).convert("RGBA")
    bbox = img.getbbox()
    if bbox:
        img = img.crop(bbox)
    img = img.resize((1180, 1180), Image.Resampling.LANCZOS)
    img = img.filter(ImageFilter.UnsharpMask(radius=2.6, percent=210, threshold=2))
    img.save(path)


def create_overview_visual(path: Path):
    img = Image.new("RGB", (1600, 900), "#F6F9FF")
    draw = ImageDraw.Draw(img)
    title_font = font(50, True)
    section_font = font(28, True)
    body_font = font(23, False)
    small_font = font(21, False)

    rounded(draw, (40, 40, 1560, 860), 36, "#FFFFFF", "#DCE7F7", 2)
    draw.text((70, 70), "Arquitetura do Sistema", fill=PRIMARY_DARK, font=title_font)
    draw.text((70, 130), "Plataforma web para operacao, controle administrativo e organizacao completa da obra.", fill=GRAY, font=body_font)

    cards = [
        ("Engenharia", "Croqui/PDF, RDO, lancamento diario e documentacao tecnica.", PRIMARY),
        ("Manutencao", "Preventiva, corretiva, alertas abertos e relatorios por equipamento.", ORANGE),
        ("Transporte", "Romaneio de carga, QR de recebimento, divergencias e geolocalizacao.", PURPLE),
        ("Almoxarifado", "Entrada central, saidas, EPI, estoque e rastreabilidade.", GREEN),
        ("Cadastros", "Obras, bases, equipamentos, funcionarios, usuarios e configuracoes.", RED),
        ("Adm. Financeiro", "Financeiro do cliente, historico, situacao do plano e acompanhamento.", PRIMARY_DARK),
    ]

    positions = [
        (70, 210, 500, 470),
        (550, 210, 980, 470),
        (1030, 210, 1490, 470),
        (70, 500, 500, 790),
        (550, 500, 980, 790),
        (1030, 500, 1490, 790),
    ]
    for (title, desc, color), pos in zip(cards, positions):
        rounded(draw, pos, 28, "#FBFDFF", color, 4)
        rounded(draw, (pos[0] + 24, pos[1] + 24, pos[0] + 88, pos[1] + 88), 18, color)
        draw.text((pos[0] + 115, pos[1] + 32), title, fill=PRIMARY_DARK, font=section_font)
        draw_wrapped(draw, desc, (pos[0] + 28, pos[1] + 112, pos[2] - pos[0] - 56, 120), fill=GRAY, font_obj=body_font, line_gap=8)
        draw.text((pos[0] + 28, pos[3] - 48), "Modulo estruturado por setor", fill=color, font=small_font)
    img.save(path)


def create_operational_visual(path: Path):
    img = Image.new("RGB", (1600, 900), "#F8FBFF")
    draw = ImageDraw.Draw(img)
    title_font = font(50, True)
    box_title = font(28, True)
    body_font = font(23, False)
    step_font = font(25, True)

    draw.text((70, 60), "Fluxo Operacional no Campo", fill=PRIMARY_DARK, font=title_font)
    draw.text((70, 122), "A operacao foi desenhada para funcionar no celular com poucos toques e foco em uso diario.", fill=GRAY, font=body_font)

    steps = [
        ("1. Acesso no celular", "Painel operacional com setores claros e modulos liberados por permissao.", PRIMARY),
        ("2. Lancamento do dia", "Diesel, RDO, croqui, manutencao, saidas e romaneio de transporte.", GREEN),
        ("3. Validacao em campo", "QR Code, assinaturas, fotos, horario e coordenadas de entrega/recebimento.", ORANGE),
        ("4. Historico e prova", "Tudo fica salvo para conferencia, relatorio e auditoria futura.", PURPLE),
    ]
    x = 80
    y = 250
    w = 330
    h = 360
    gap = 40
    for idx, (title, desc, color) in enumerate(steps):
        rounded(draw, (x, y, x + w, y + h), 28, "#FFFFFF", color, 4)
        rounded(draw, (x + 24, y + 24, x + 82, y + 82), 16, color)
        draw.text((x + 44, y + 34), str(idx + 1), fill="white", font=step_font)
        draw_wrapped(draw, title, (x + 24, y + 110, w - 48, 60), fill=PRIMARY_DARK, font_obj=box_title, line_gap=6)
        draw_wrapped(draw, desc, (x + 24, y + 190, w - 48, 140), fill=GRAY, font_obj=body_font, line_gap=8)
        if idx < len(steps) - 1:
            mid_y = y + h / 2
            draw.line((x + w + 10, mid_y, x + w + gap - 10, mid_y), fill=color, width=8)
            draw.polygon(
                [(x + w + gap - 12, mid_y), (x + w + gap - 34, mid_y - 14), (x + w + gap - 34, mid_y + 14)],
                fill=color,
            )
        x += w + gap
    img.save(path)


def create_admin_visual(path: Path):
    img = Image.new("RGB", (1600, 900), "#F7F9FE")
    draw = ImageDraw.Draw(img)
    title_font = font(48, True)
    section_font = font(28, True)
    body_font = font(23, False)
    small_font = font(21, False)

    draw.text((70, 60), "Painel Administrativo do Cliente", fill=PRIMARY_DARK, font=title_font)
    draw.text((70, 118), "Organizacao dos cadastros, controles, relatorios e rotinas administrativas do dia a dia.", fill=GRAY, font=body_font)

    rounded(draw, (70, 200, 760, 780), 30, "#FFFFFF", "#DCE7F7", 2)
    draw.text((100, 235), "Cadastros e Controle", fill=PRIMARY_DARK, font=section_font)
    items_left = [
        ("Cadastros", "Obras, bases, equipamentos, funcionarios e usuarios."),
        ("Controle", "Transferencias, transportes, manutencao e almoxarifado."),
        ("Relatorios", "RDO, croqui, diesel, manutencao, transferencias e transportes."),
        ("Permissoes", "Acesso por perfil, base e cidade, com visao operacional e administrativa."),
    ]
    top = 300
    for title, desc in items_left:
        rounded(draw, (100, top, 720, top + 96), 22, PRIMARY_SOFT, "#C9D9FF", 2)
        draw.text((126, top + 18), title, fill=PRIMARY_DARK, font=section_font)
        draw.text((126, top + 54), desc, fill=GRAY, font=small_font)
        top += 114

    rounded(draw, (840, 200, 1530, 780), 30, "#FFFFFF", "#DCE7F7", 2)
    draw.text((870, 235), "Financeiro e Governanca", fill=PRIMARY_DARK, font=section_font)
    items_right = [
        ("Financeiro", "Plano, valor, vencimento e visao dos pagamentos do cliente."),
        ("Bloqueios e liberacoes", "Regras para manter o sistema funcionando de forma organizada e segura."),
        ("Relatorios", "Historico e acompanhamento gerencial das movimentacoes administrativas."),
        ("Escalabilidade", "Estrutura preparada para crescer junto com novas obras, bases e equipes."),
    ]
    top = 300
    for title, desc in items_right:
        rounded(draw, (870, top, 1490, top + 96), 22, "#FFF6EA", "#FFD9A6", 2)
        draw.text((896, top + 18), title, fill=PRIMARY_DARK, font=section_font)
        draw.text((896, top + 54), desc, fill=GRAY, font=small_font)
        top += 114
    img.save(path)


def create_mobile_visual(path: Path):
    img = Image.new("RGB", (1600, 900), "#EEF4FF")
    draw = ImageDraw.Draw(img)
    title_font = font(48, True)
    body_font = font(23, False)
    card_title = font(26, True)

    draw.text((70, 60), "Experiencia Mobile / App Instalavel", fill=PRIMARY_DARK, font=title_font)
    draw.text((70, 118), "A aplicacao pode ser instalada no Android como app e usada em campo com foco em velocidade e clareza.", fill=GRAY, font=body_font)

    phone = (140, 180, 640, 820)
    rounded(draw, phone, 46, "#0B1D35")
    rounded(draw, (165, 215, 615, 795), 32, "#FFFFFF")
    rounded(draw, (265, 240, 515, 268), 14, "#D9E7FF")
    rounded(draw, (205, 300, 575, 430), 28, "#0B5ED7")
    draw.text((244, 332), "Setores e modulos", fill="white", font=card_title)
    draw.text((244, 372), "Painel\nEngenharia\nManutencao\nTransporte", fill="white", font=body_font)
    rounded(draw, (205, 460, 575, 720), 28, "#F5F8FF", "#D7E4FF", 2)
    draw.text((232, 492), "Uso de campo", fill=PRIMARY_DARK, font=card_title)
    draw.text((232, 532), "- lancamentos diarios\n- QR code\n- assinatura\n- localizacao\n- consulta rapida", fill=GRAY, font=body_font)

    notes = [
        ("Instalacao simples", "Abrir no navegador e adicionar a tela inicial como aplicativo."),
        ("Foco operacional", "Modulos organizados para reduzir poluicao visual e acelerar o uso no celular."),
        ("Rastreabilidade", "QR + assinatura + localizacao deixam o historico pronto para conferencia."),
    ]
    top = 220
    for idx, (title, desc) in enumerate(notes):
        y = top + idx * 180
        rounded(draw, (760, y, 1490, y + 140), 28, "#FFFFFF", "#DCE7F7", 2)
        draw.text((792, y + 24), title, fill=PRIMARY_DARK, font=card_title)
        draw_wrapped(draw, desc, (792, y + 62, 640, 56), fill=GRAY, font_obj=body_font, line_gap=6)
    img.save(path)


def screenshot_slots():
    return [
        ("01_painel.png", "Painel principal"),
        ("02_lancamento_diario.png", "Lançamento diário"),
        ("03_transporte.png", "Romaneio de transporte"),
        ("04_almoxarifado.png", "Almoxarifado / entrada"),
    ]


def create_screenshots_board(path: Path):
    img = Image.new("RGB", (1600, 980), "#F7FAFF")
    draw = ImageDraw.Draw(img)
    title_font = font(48, True)
    body_font = font(22, False)
    label_font = font(24, True)
    note_font = font(19, False)

    draw.text((70, 56), "Telas do Sistema", fill=PRIMARY_DARK, font=title_font)
    draw.text((70, 116), "Esta pagina foi preparada para receber capturas reais das telas mais importantes do produto.", fill=GRAY, font=body_font)

    slots = screenshot_slots()
    positions = [
        (70, 190, 760, 530),
        (840, 190, 1530, 530),
        (70, 580, 760, 920),
        (840, 580, 1530, 920),
    ]

    for (filename, label), pos in zip(slots, positions):
        rounded(draw, pos, 28, "#FFFFFF", "#DCE7F7", 2)
        target = SCREENSHOT_DIR / filename
        image_box = (pos[0] + 18, pos[1] + 18, pos[2] - 18, pos[1] + 232)
        if target.exists():
            shot = Image.open(target).convert("RGB")
            shot.thumbnail((image_box[2] - image_box[0], image_box[3] - image_box[1]), Image.Resampling.LANCZOS)
            paste_x = image_box[0] + ((image_box[2] - image_box[0]) - shot.width) // 2
            paste_y = image_box[1] + ((image_box[3] - image_box[1]) - shot.height) // 2
            img.paste(shot, (paste_x, paste_y))
        else:
            rounded(draw, image_box, 18, "#EEF4FF", "#CFE0FF", 2)
            draw.text((image_box[0] + 24, image_box[1] + 70), "Adicione aqui o print real", fill=PRIMARY, font=label_font)
            draw.text((image_box[0] + 24, image_box[1] + 112), filename, fill=GRAY, font=note_font)
        draw.text((pos[0] + 24, pos[1] + 254), label, fill=PRIMARY_DARK, font=label_font)
        draw_wrapped(
            draw,
            "Quando o print for colocado na pasta 'portfolio/prints', ele entra automaticamente nesta pagina na proxima geracao do arquivo.",
            (pos[0] + 24, pos[1] + 288, pos[2] - pos[0] - 48, 90),
            fill=GRAY,
            font_obj=note_font,
            line_gap=6,
        )
    img.save(path)


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_default_font(run, size, bold=False, color=PRIMARY_DARK):
    run.font.name = "Arial"
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.color.rgb = RGBColor(*hex_rgb(color))
    run._element.rPr.rFonts.set(qn("w:ascii"), "Arial")
    run._element.rPr.rFonts.set(qn("w:hAnsi"), "Arial")


def add_styled_paragraph(doc, text, style_name=None, size=12, bold=False, color=PRIMARY_DARK, align=WD_ALIGN_PARAGRAPH.LEFT, space_after=6):
    p = doc.add_paragraph(style=style_name)
    p.alignment = align
    run = p.add_run(text)
    set_default_font(run, size=size, bold=bold, color=color)
    p.paragraph_format.space_after = Pt(space_after)
    p.paragraph_format.line_spacing = 1.08
    return p


def add_bullet(doc, text, color=GRAY):
    p = doc.add_paragraph(style="List Bullet")
    run = p.add_run(text)
    set_default_font(run, size=11, color=color)
    p.paragraph_format.space_after = Pt(6)
    return p


def add_cover_callout(doc, title, lines):
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    cell = table.rows[0].cells[0]
    set_cell_shading(cell, "EEF4FF")
    p_title = cell.paragraphs[0]
    p_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_title = p_title.add_run(title)
    set_default_font(run_title, size=13, bold=True, color=PRIMARY_DARK)
    p_title.paragraph_format.space_after = Pt(5)
    for idx, line in enumerate(lines):
      p = cell.add_paragraph()
      p.alignment = WD_ALIGN_PARAGRAPH.CENTER
      if idx == 0:
          p.paragraph_format.space_before = Pt(0)
      p.paragraph_format.space_after = Pt(3)
      r = p.add_run(line)
      set_default_font(r, size=10, color=GRAY)
    return table


def add_cover_highlights(doc, items):
    table = doc.add_table(rows=1, cols=len(items))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    for idx, (title, desc, fill) in enumerate(items):
        cell = table.rows[0].cells[idx]
        set_cell_shading(cell, fill)
        p1 = cell.paragraphs[0]
        p1.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r1 = p1.add_run(title)
        set_default_font(r1, size=12, bold=True, color=PRIMARY_DARK)
        p1.paragraph_format.space_after = Pt(4)
        p2 = cell.add_paragraph()
        p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p2.paragraph_format.space_after = Pt(2)
        r2 = p2.add_run(desc)
        set_default_font(r2, size=9, color=GRAY)
    return table


def create_doc(assets):
    doc = Document()
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.left_margin = Inches(0.55)
    section.right_margin = Inches(0.55)
    section.top_margin = Inches(0.55)
    section.bottom_margin = Inches(0.55)

    styles = doc.styles
    styles["Normal"].font.name = "Arial"
    styles["Normal"].font.size = Pt(11)

    header = section.header
    hp = header.paragraphs[0]
    hp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    hr = hp.add_run("Equipamento Gestao | Portfolio do Sistema")
    set_default_font(hr, 9, color=GRAY)

    logo_paragraph = doc.add_paragraph()
    logo_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    logo_run = logo_paragraph.add_run()
    logo_run.add_picture(str(COVER_LOGO_PATH), width=Inches(3.15))
    logo_paragraph.paragraph_format.space_after = Pt(10)
    add_styled_paragraph(doc, "PORTFOLIO COMERCIAL DO SISTEMA", size=11, bold=True, color=GRAY, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=2)
    add_styled_paragraph(doc, "Equipamento Gestao", size=32, bold=True, color=PRIMARY_DARK, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=1)
    add_styled_paragraph(doc, "Controle operacional e administrativo para obras", size=18, bold=True, color=PRIMARY, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=4)
    add_styled_paragraph(
        doc,
        "Uma plataforma pensada para centralizar a rotina da obra, reduzir falhas operacionais e dar visibilidade real ao que acontece no campo e no escritorio.",
        size=11,
        color=GRAY,
        align=WD_ALIGN_PARAGRAPH.CENTER,
        space_after=8,
    )
    add_cover_highlights(
        doc,
        [
            ("Operacao em campo", "Lancamentos, RDO, QR, assinatura e uso no celular.", "EAF2FF"),
            ("Controle interno", "Cadastros, manutencao, transporte e almoxarifado.", "EEF9F2"),
            ("Visao gerencial", "Historico, financeiro do cliente e governanca.", "FFF4E6"),
        ],
    )
    doc.add_paragraph().paragraph_format.space_after = Pt(6)
    add_cover_callout(
        doc,
        "Proposta de valor",
        [
            "Sistema responsivo para computador e celular",
            "Modulos organizados por setor com foco em uso diario",
            "Mais controle, mais rastreabilidade e menos planilhas soltas",
        ],
    )
    doc.add_paragraph().paragraph_format.space_after = Pt(8)
    doc.add_picture(str(assets["overview"]), width=IMAGE_WIDTH)

    add_styled_paragraph(doc, "Resumo executivo", style_name="Heading 1", size=16, bold=True, color=PRIMARY_DARK, space_after=4)
    for bullet in [
        "Sistema web responsivo para operacao diaria no computador e no celular.",
        "Estrutura por setores: Engenharia, Manutencao, Transporte, Almoxarifado, Cadastros e Administrativo Financeiro.",
        "Financeiro do cliente com visao de plano, vencimento, historico e governanca.",
        "Permissoes por perfil, base, cidade e visao administrativa ou operacional.",
    ]:
        add_bullet(doc, bullet)

    doc.add_page_break()
    add_styled_paragraph(doc, "1. Visao geral da plataforma", style_name="Heading 1", size=18, bold=True, color=PRIMARY_DARK, space_after=4)
    add_styled_paragraph(
        doc,
        "O sistema foi desenhado para concentrar toda a rotina da obra em uma unica plataforma: lancamentos do dia, controle tecnico, manutencao, transporte, almoxarifado e gestao administrativa.",
        size=11,
        color=GRAY,
        space_after=10,
    )
    doc.add_picture(str(assets["overview"]), width=IMAGE_WIDTH)

    add_styled_paragraph(doc, "Principais blocos do sistema", style_name="Heading 2", size=14, bold=True, color=PRIMARY_DARK, space_after=3)
    table = doc.add_table(rows=1, cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    hdr = table.rows[0].cells
    hdr[0].text = "Area"
    hdr[1].text = "O que entrega"
    set_cell_shading(hdr[0], "EAF2FF")
    set_cell_shading(hdr[1], "EAF2FF")
    rows = [
        ("Operacional", "Lancamento diario, RDO, croqui/PDF, transporte, QR, assinatura, manutencao e saídas."),
        ("Administrativo", "Cadastros, controle de equipamentos, funcionarios, bases, obras e configuracoes."),
        ("Administrativo Financeiro", "Plano, historico, acompanhamento financeiro e governanca do cliente."),
    ]
    for left, right in rows:
        cells = table.add_row().cells
        cells[0].text = left
        cells[1].text = right

    doc.add_page_break()
    add_styled_paragraph(doc, "2. Operacao no campo", style_name="Heading 1", size=18, bold=True, color=PRIMARY_DARK, space_after=4)
    add_styled_paragraph(
        doc,
        "No uso operacional, o foco e velocidade. A aplicacao no celular concentra somente os modulos que precisam rodar no dia, reduzindo poluicao visual e acelerando a execucao em campo.",
        size=11,
        color=GRAY,
        space_after=10,
    )
    doc.add_picture(str(assets["operational"]), width=IMAGE_WIDTH)
    for bullet in [
        "Lançamento diário de equipamentos e produção.",
        "RDO e produção de campo com croqui/PDF.",
        "Abastecimento, manutenção, saídas de materiais e EPI.",
        "Romaneio de transporte com QR Code, assinatura e geolocalização.",
    ]:
        add_bullet(doc, bullet)

    doc.add_page_break()
    add_styled_paragraph(doc, "3. Painel administrativo do cliente", style_name="Heading 1", size=18, bold=True, color=PRIMARY_DARK, space_after=4)
    add_styled_paragraph(
        doc,
        "A camada administrativa organiza os cadastros, relatorios, controles e acompanhamento financeiro do proprio cliente, sem depender de planilhas ou registros paralelos.",
        size=11,
        color=GRAY,
        space_after=10,
    )
    doc.add_picture(str(assets["admin"]), width=IMAGE_WIDTH)
    for bullet in [
        "Cadastros estruturais de obras, bases, equipamentos, funcionários e usuários operacionais.",
        "Relatórios administrativos para transporte, transferências, diesel, manutenção e produção.",
        "Financeiro do cliente com visao de plano, vencimento, historico e liberacoes.",
        "Organizacao administrativa para manter operacao, controle e auditoria no mesmo ambiente.",
    ]:
        add_bullet(doc, bullet)

    doc.add_page_break()
    add_styled_paragraph(doc, "4. Mobilidade e experiencia de uso", style_name="Heading 1", size=18, bold=True, color=PRIMARY_DARK, space_after=4)
    add_styled_paragraph(
        doc,
        "A plataforma pode ser instalada no Android como aplicativo, com navegação simplificada, ícone próprio e experiência focada em operação rápida.",
        size=11,
        color=GRAY,
        space_after=10,
    )
    doc.add_picture(str(assets["mobile"]), width=IMAGE_WIDTH)
    for bullet in [
        "Instalação como app pela tela inicial do celular.",
        "Logo personalizada, experiência PWA e tela inicial mais profissional.",
        "Setores visíveis sem arrastar para o lado e módulos organizados por contexto.",
    ]:
        add_bullet(doc, bullet)

    doc.add_page_break()
    add_styled_paragraph(doc, "5. Diferenciais competitivos", style_name="Heading 1", size=18, bold=True, color=PRIMARY_DARK, space_after=6)
    for bullet in [
        "Sistema pensado para a rotina real de obra, não apenas para escritório.",
        "Rastreabilidade operacional com QR Code, assinatura, horário e localização.",
        "Separação por perfil, base e cidade para reduzir erro e dar governança.",
        "Estrutura pronta para seguir evoluindo com novos módulos sem desmontar o sistema.",
    ]:
        add_bullet(doc, bullet)

    add_styled_paragraph(doc, "Observacao", style_name="Heading 2", size=14, bold=True, color=PRIMARY_DARK, space_after=2)
    add_styled_paragraph(
        doc,
        "Esta e uma primeira versao institucional do portfolio. Na proxima etapa, as imagens graficas podem ser complementadas por capturas reais das telas principais do sistema para reforco comercial.",
        size=10,
        color=GRAY,
        space_after=4,
    )

    doc.add_page_break()
    add_styled_paragraph(doc, "6. Galeria de telas reais", style_name="Heading 1", size=18, bold=True, color=PRIMARY_DARK, space_after=4)
    add_styled_paragraph(
        doc,
        "Esta pagina foi preparada para receber capturas reais das telas do sistema. Basta salvar os prints na pasta 'portfolio/prints' com os nomes definidos no roteiro e gerar novamente o documento.",
        size=11,
        color=GRAY,
        space_after=10,
    )
    doc.add_picture(str(assets["screenshots"]), width=IMAGE_WIDTH)

    footer = section.footer
    fp = footer.paragraphs[0]
    fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    fr = fp.add_run("Equipamento Gestao | Controle operacional e administrativo")
    set_default_font(fr, 9, color=GRAY)

    doc.save(DOCX_PATH)


def main():
    OUT_DIR.mkdir(parents=True, exist_ok=True)
    assets = {
        "overview": OUT_DIR / "portfolio_overview.png",
        "operational": OUT_DIR / "portfolio_operational.png",
        "admin": OUT_DIR / "portfolio_admin.png",
        "mobile": OUT_DIR / "portfolio_mobile.png",
        "screenshots": OUT_DIR / "portfolio_screenshots.png",
    }
    SCREENSHOT_DIR.mkdir(parents=True, exist_ok=True)
    create_cover_logo(COVER_LOGO_PATH)
    create_overview_visual(assets["overview"])
    create_operational_visual(assets["operational"])
    create_admin_visual(assets["admin"])
    create_mobile_visual(assets["mobile"])
    create_screenshots_board(assets["screenshots"])
    create_doc(assets)
    print(DOCX_PATH)


if __name__ == "__main__":
    main()
