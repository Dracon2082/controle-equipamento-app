from pathlib import Path

from PIL import Image, ImageDraw, ImageFont
from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(r"C:\Users\amare\controle-equipamento")
OUT_DIR = ROOT / "portfolio"
DOCX_PATH = OUT_DIR / "Portfolio_Planos_Equipamento_Gestao_v1.docx"
HERO_IMAGE = OUT_DIR / "pricing_cover.png"
PLANS_IMAGE = OUT_DIR / "pricing_cards.png"
LOGO_PATH = ROOT / "app" / "src" / "assets" / "logo-sistema.png"

PRIMARY = "#0B5ED7"
PRIMARY_DARK = "#163256"
PRIMARY_SOFT = "#EAF2FF"
GREEN_SOFT = "#EAF9EF"
ORANGE_SOFT = "#FFF4E6"
PURPLE_SOFT = "#F4EEFF"
GRAY = "#5F6F86"
BORDER = "#D9E4F2"

PLANS = [
    {"nome": "Plano 1", "preco": 349, "gestores": 1, "admins": 7, "operadores": "30", "ideal": "Operacao pequena ou media"},
    {"nome": "Plano 2", "preco": 499, "gestores": 1, "admins": 10, "operadores": "50", "ideal": "Empresa em crescimento"},
    {"nome": "Plano 3", "preco": 699, "gestores": 1, "admins": 15, "operadores": "80", "ideal": "Operacao com varias frentes"},
    {"nome": "Plano 4", "preco": 899, "gestores": 1, "admins": 20, "operadores": "Ilimitados", "ideal": "Operacao grande e escalavel"},
]

CYCLES = [
    {"nome": "Mensal", "meses": 1, "desconto": 0},
    {"nome": "3 meses", "meses": 3, "desconto": 5},
    {"nome": "6 meses", "meses": 6, "desconto": 10},
    {"nome": "12 meses", "meses": 12, "desconto": 15},
]


def hex_rgb(value: str):
    value = value.lstrip("#")
    return tuple(int(value[i : i + 2], 16) for i in (0, 2, 4))


def money(value):
    return f"R$ {value:,.2f}".replace(",", "X").replace(".", ",").replace("X", ".")


def font(size: int, bold: bool = False):
    candidates = [
        ("arialbd.ttf" if bold else "arial.ttf"),
        ("C:/Windows/Fonts/arialbd.ttf" if bold else "C:/Windows/Fonts/arial.ttf"),
        ("C:/Windows/Fonts/calibrib.ttf" if bold else "C:/Windows/Fonts/calibri.ttf"),
    ]
    for candidate in candidates:
        try:
            return ImageFont.truetype(candidate, size=size)
        except OSError:
            continue
    return ImageFont.load_default()


def rounded(draw, xy, radius, fill, outline=None, width=1):
    draw.rounded_rectangle(xy, radius=radius, fill=fill, outline=outline, width=width)


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_default_font(run, size, bold=False, color=PRIMARY_DARK):
    run.font.name = "Arial"
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.color.rgb = RGBColor(*hex_rgb(color))
    run._element.rPr.rFonts.set(qn("w:ascii"), "Arial")
    run._element.rPr.rFonts.set(qn("w:hAnsi"), "Arial")


def add_styled_paragraph(doc, text, size=12, bold=False, color=PRIMARY_DARK, align=WD_ALIGN_PARAGRAPH.LEFT, space_after=6):
    p = doc.add_paragraph()
    p.alignment = align
    r = p.add_run(text)
    set_default_font(r, size=size, bold=bold, color=color)
    p.paragraph_format.space_after = Pt(space_after)
    p.paragraph_format.line_spacing = 1.08
    return p


def add_bullet(doc, text, color=GRAY):
    p = doc.add_paragraph(style="List Bullet")
    r = p.add_run(text)
    set_default_font(r, size=11, color=color)
    p.paragraph_format.space_after = Pt(6)
    return p


def calc_cycle(preco, meses, desconto_pct):
    bruto = preco * meses
    desconto = round(bruto * desconto_pct / 100, 2)
    total = round(bruto - desconto, 2)
    mensal_eq = round(total / meses, 2)
    return bruto, desconto, total, mensal_eq


def create_cover_image():
    img = Image.new("RGB", (1800, 940), "#F5F8FE")
    draw = ImageDraw.Draw(img)
    rounded(draw, (40, 40, 1760, 900), 38, fill="#FFFFFF", outline="#DCE6F3", width=2)
    rounded(draw, (90, 110, 810, 820), 34, fill=PRIMARY, outline=PRIMARY, width=2)
    rounded(draw, (920, 110, 1690, 360), 34, fill=PRIMARY_SOFT, outline="#CFE0FF", width=2)
    rounded(draw, (920, 395, 1690, 820), 34, fill="#FFFFFF", outline="#E1EAF6", width=2)

    if LOGO_PATH.exists():
        logo = Image.open(LOGO_PATH).convert("RGBA")
        logo.thumbnail((260, 260))
        img.paste(logo, (310 - logo.width // 2, 155), logo)

    draw.text((120, 400), "Planos Comerciais", fill="#FFFFFF", font=font(58, True))
    draw.text((120, 480), "Equipamento Gestao", fill="#FFFFFF", font=font(40, True))
    draw.text((120, 555), "Controle operacional e administrativo\npara obras, transportes e almoxarifado.", fill="#EAF3FF", font=font(25))
    draw.text((120, 705), "Mais previsibilidade, mais controle\n e uma estrutura de usuarios escalavel.", fill="#DDE9FF", font=font(22))

    draw.text((970, 150), "Estrutura pronta para vender", fill=PRIMARY_DARK, font=font(32, True))
    draw.text((970, 205), "Planos mensais e pacotes com fidelizacao,\ncom escalada clara por tamanho de operacao.", fill=GRAY, font=font(22))

    benefits = [
        ("1 gestor", "Governanca central da operacao"),
        ("ADMs por faixa", "Mais controle para escritorio e bases"),
        ("Operadores escalaveis", "Do campo ate a operacao ampliada"),
    ]
    y = 430
    for title, desc in benefits:
        rounded(draw, (970, y, 1640, y + 90), 24, fill="#F9FBFF", outline="#E1EAF6", width=2)
        draw.text((1000, y + 16), title, fill=PRIMARY, font=font(24, True))
        draw.text((1000, y + 50), desc, fill=GRAY, font=font(20))
        y += 108

    img.save(HERO_IMAGE)


def create_plans_image():
    img = Image.new("RGB", (1800, 1180), "#F5F8FE")
    draw = ImageDraw.Draw(img)
    colors = [PRIMARY_SOFT, GREEN_SOFT, ORANGE_SOFT, PURPLE_SOFT]
    accent = [PRIMARY, "#198754", "#E67700", "#6F42C1"]

    card_w, card_h = 790, 250
    positions = [(70, 80), (940, 80), (70, 380), (940, 380)]

    for idx, plan in enumerate(PLANS):
        x, y = positions[idx]
        rounded(draw, (x, y, x + card_w, y + card_h), 30, fill=colors[idx], outline="#D7E4F5", width=2)
        draw.rounded_rectangle((x + 16, y + 18, x + 150, y + 68), radius=22, fill=accent[idx])
        draw.text((x + 44, y + 30), plan["nome"], fill="#FFFFFF", font=font(22, True))
        draw.text((x + 38, y + 92), money(plan["preco"]), fill=PRIMARY_DARK, font=font(44, True))
        draw.text((x + 38, y + 148), f'{plan["gestores"]} gestor | {plan["admins"]} ADMs | {plan["operadores"]} operadores', fill=PRIMARY_DARK, font=font(24, True))
        draw.text((x + 38, y + 192), plan["ideal"], fill=GRAY, font=font(20))

    rounded(draw, (70, 710, 1730, 1080), 30, fill="#FFFFFF", outline="#D7E4F5", width=2)
    draw.text((110, 750), "Pacotes com fidelizacao", fill=PRIMARY_DARK, font=font(34, True))
    draw.text((110, 805), "Descontos automaticos para contratos de maior prazo.", fill=GRAY, font=font(22))

    box_x = 110
    for cycle in CYCLES:
        bruto, desconto, total, mensal_eq = calc_cycle(PLANS[0]["preco"], cycle["meses"], cycle["desconto"])
        rounded(draw, (box_x, 860, box_x + 360, 1015), 24, fill="#F7FAFF", outline="#DFE8F5", width=2)
        draw.text((box_x + 26, 885), cycle["nome"], fill=PRIMARY, font=font(24, True))
        if cycle["meses"] == 1:
            draw.text((box_x + 26, 930), f"Sem desconto | {money(PLANS[0]['preco'])}/mes", fill=PRIMARY_DARK, font=font(20, True))
        else:
            draw.text((box_x + 26, 930), f"{cycle['desconto']}% off | Total {money(total)}", fill=PRIMARY_DARK, font=font(20, True))
            draw.text((box_x + 26, 968), f"Equivale a {money(mensal_eq)}/mes", fill=GRAY, font=font(18))
        box_x += 400

    img.save(PLANS_IMAGE)


def create_doc():
    doc = Document()
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.left_margin = Inches(0.55)
    section.right_margin = Inches(0.55)
    section.top_margin = Inches(0.55)
    section.bottom_margin = Inches(0.55)

    styles = doc.styles
    styles["Normal"].font.name = "Arial"
    styles["Normal"].font.size = Pt(11)

    header = section.header
    hp = header.paragraphs[0]
    hp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    hr = hp.add_run("Equipamento Gestao | Tabela comercial de planos")
    set_default_font(hr, 9, color=GRAY)

    doc.add_picture(str(HERO_IMAGE), width=Inches(7.3))
    add_styled_paragraph(doc, "Apresentacao comercial dos planos", size=11, bold=True, color=GRAY, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=2)
    add_styled_paragraph(doc, "Planos, usuarios e fidelizacao", size=26, bold=True, color=PRIMARY_DARK, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=4)
    add_styled_paragraph(
        doc,
        "Uma estrutura clara para vender o sistema por porte de operacao, com desconto progressivo para contratos de maior prazo.",
        size=11,
        color=GRAY,
        align=WD_ALIGN_PARAGRAPH.CENTER,
        space_after=10,
    )

    highlight = doc.add_table(rows=1, cols=3)
    highlight.alignment = WD_TABLE_ALIGNMENT.CENTER
    highlight.style = "Table Grid"
    items = [
        ("Base mensal", money(349), "Entrada comercial forte para pequenas e medias operacoes.", "EAF2FF"),
        ("Fidelizacao", "3, 6 e 12 meses", "Desconto progressivo para aumentar recorrencia e previsibilidade.", "EEF9F2"),
        ("Escalabilidade", "Ate ilimitado", "Do controle inicial ate operacao grande com varias equipes.", "FFF4E6"),
    ]
    for idx, (title, value, desc, fill) in enumerate(items):
        cell = highlight.rows[0].cells[idx]
        set_cell_shading(cell, fill)
        p1 = cell.paragraphs[0]
        p1.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r1 = p1.add_run(title)
        set_default_font(r1, 12, True)
        p2 = cell.add_paragraph()
        p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r2 = p2.add_run(value)
        set_default_font(r2, 16, True, PRIMARY)
        p3 = cell.add_paragraph()
        p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r3 = p3.add_run(desc)
        set_default_font(r3, 9, False, GRAY)

    doc.add_page_break()
    add_styled_paragraph(doc, "1. Visao geral dos planos", size=18, bold=True, color=PRIMARY_DARK, space_after=4)
    add_styled_paragraph(doc, "Os planos foram organizados para acompanhar o crescimento da operacao, mantendo uma leitura comercial simples: mais usuarios, mais estrutura e mais capacidade de escala.", size=11, color=GRAY, space_after=8)
    doc.add_picture(str(PLANS_IMAGE), width=Inches(7.3))

    add_styled_paragraph(doc, "2. Tabela comparativa", size=18, bold=True, color=PRIMARY_DARK, space_after=4)
    table = doc.add_table(rows=1, cols=5)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    headers = ["Plano", "Valor mensal", "Gestores / ADMs", "Operadores", "Perfil ideal"]
    for idx, title in enumerate(headers):
        cell = table.rows[0].cells[idx]
        cell.text = title
        set_cell_shading(cell, "EAF2FF")

    for plan in PLANS:
        row = table.add_row().cells
        row[0].text = plan["nome"]
        row[1].text = money(plan["preco"])
        row[2].text = f'{plan["gestores"]} gestor | {plan["admins"]} ADMs'
        row[3].text = str(plan["operadores"])
        row[4].text = plan["ideal"]

    add_styled_paragraph(doc, "3. Fidelizacao e descontos", size=18, bold=True, color=PRIMARY_DARK, space_after=4)
    add_styled_paragraph(doc, "Os pacotes de maior prazo ajudam a tornar a venda mais atrativa sem desvalorizar o produto. Abaixo, um exemplo tomando como base o Plano 1.", size=11, color=GRAY, space_after=8)

    cycle_table = doc.add_table(rows=1, cols=5)
    cycle_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    cycle_table.style = "Table Grid"
    cycle_headers = ["Ciclo", "Desconto", "Valor bruto", "Total com desconto", "Equivalente mensal"]
    for idx, title in enumerate(cycle_headers):
        cell = cycle_table.rows[0].cells[idx]
        cell.text = title
        set_cell_shading(cell, "EEF9F2")

    for cycle in CYCLES:
        bruto, desconto, total, mensal_eq = calc_cycle(PLANS[0]["preco"], cycle["meses"], cycle["desconto"])
        row = cycle_table.add_row().cells
        row[0].text = cycle["nome"]
        row[1].text = f'{cycle["desconto"]}%'
        row[2].text = money(bruto)
        row[3].text = money(total)
        row[4].text = money(mensal_eq) if cycle["meses"] > 1 else money(PLANS[0]["preco"])

    add_styled_paragraph(doc, "4. Como vender cada faixa", size=18, bold=True, color=PRIMARY_DARK, space_after=4)
    for bullet in [
        "Plano 1: ideal para entrar em empresas pequenas ou medias, com estrutura completa e preco forte de entrada.",
        "Plano 2: indicado para cliente em crescimento, com mais usuarios administrativos e mais operadores em campo.",
        "Plano 3: direcionado para operacoes maiores, com varias frentes e maior volume de controle diario.",
        "Plano 4: opcao premium para empresas maiores, com 20 ADMs e operadores ilimitados.",
        "Pacotes de 6 e 12 meses devem ser oferecidos como alternativa para aumentar previsibilidade de receita.",
    ]:
        add_bullet(doc, bullet)

    add_styled_paragraph(doc, "5. Observacao comercial", size=18, bold=True, color=PRIMARY_DARK, space_after=4)
    add_styled_paragraph(
        doc,
        "Os valores podem ser usados como base oficial da apresentacao comercial. Em negociacoes especificas, e possivel conceder ajuste pontual sem descaracterizar a escada principal de precos.",
        size=11,
        color=GRAY,
        space_after=8,
    )

    footer = section.footer
    fp = footer.paragraphs[0]
    fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    fr = fp.add_run("Equipamento Gestao | Planos comerciais e fidelizacao")
    set_default_font(fr, 9, color=GRAY)

    doc.save(DOCX_PATH)


def main():
    OUT_DIR.mkdir(parents=True, exist_ok=True)
    create_cover_image()
    create_plans_image()
    create_doc()
    print(DOCX_PATH)


if __name__ == "__main__":
    main()
