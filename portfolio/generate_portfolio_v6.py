from pathlib import Path

from PIL import Image, ImageDraw, ImageFont
from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(r"C:\Users\amare\controle-equipamento")
OUT_DIR = ROOT / "portfolio"
DOCX_PATH = OUT_DIR / "Portfolio_Equipamento_Gestao_v6.docx"
LOGO_PATH = ROOT / "app" / "src" / "assets" / "logo-sistema.png"
SCREENSHOT_DIR = OUT_DIR / "prints"
PROCESSED_DIR = OUT_DIR / "prints_processados"
OVERVIEW_PATH = OUT_DIR / "portfolio_overview.png"
OPERATIONAL_PATH = OUT_DIR / "portfolio_operational.png"
ADMIN_PATH = OUT_DIR / "portfolio_admin.png"
MOBILE_PATH = OUT_DIR / "portfolio_mobile.png"
COVER_LOGO_PATH = OUT_DIR / "portfolio_cover_logo.png"
IMAGE_WIDTH = Inches(7.05)

PRIMARY = "#0B5ED7"
PRIMARY_DARK = "#163256"
GRAY = "#5F6F86"


def hex_rgb(value: str):
    value = value.lstrip("#")
    return tuple(int(value[i : i + 2], 16) for i in (0, 2, 4))


def font(size: int, bold: bool = False):
    candidates = [
        ("arialbd.ttf" if bold else "arial.ttf"),
        ("C:/Windows/Fonts/arialbd.ttf" if bold else "C:/Windows/Fonts/arial.ttf"),
        ("C:/Windows/Fonts/calibrib.ttf" if bold else "C:/Windows/Fonts/calibri.ttf"),
    ]
    for candidate in candidates:
        try:
            return ImageFont.truetype(candidate, size=size)
        except OSError:
            continue
    return ImageFont.load_default()


def rounded(draw, xy, radius, fill, outline=None, width=1):
    draw.rounded_rectangle(xy, radius=radius, fill=fill, outline=outline, width=width)


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_default_font(run, size, bold=False, color=PRIMARY_DARK):
    run.font.name = "Arial"
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.color.rgb = RGBColor(*hex_rgb(color))
    run._element.rPr.rFonts.set(qn("w:ascii"), "Arial")
    run._element.rPr.rFonts.set(qn("w:hAnsi"), "Arial")


def add_styled_paragraph(doc, text, style_name=None, size=12, bold=False, color=PRIMARY_DARK, align=WD_ALIGN_PARAGRAPH.LEFT, space_after=6):
    p = doc.add_paragraph(style=style_name)
    p.alignment = align
    run = p.add_run(text)
    set_default_font(run, size=size, bold=bold, color=color)
    p.paragraph_format.space_after = Pt(space_after)
    p.paragraph_format.line_spacing = 1.08
    return p


def add_bullet(doc, text, color=GRAY):
    p = doc.add_paragraph(style="List Bullet")
    run = p.add_run(text)
    set_default_font(run, size=11, color=color)
    p.paragraph_format.space_after = Pt(6)
    return p


def add_cover_highlights(doc):
    items = [
        ("Operacao em campo", "Lancamentos, RDO, QR, assinatura e uso no celular.", "EAF2FF"),
        ("Controle interno", "Cadastros, manutencao, transporte e almoxarifado.", "EEF9F2"),
        ("Visao gerencial", "Historico, financeiro do cliente e governanca.", "FFF4E6"),
    ]
    table = doc.add_table(rows=1, cols=len(items))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    for idx, (title, desc, fill) in enumerate(items):
        cell = table.rows[0].cells[idx]
        set_cell_shading(cell, fill)
        p1 = cell.paragraphs[0]
        p1.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r1 = p1.add_run(title)
        set_default_font(r1, size=12, bold=True)
        p1.paragraph_format.space_after = Pt(4)
        p2 = cell.add_paragraph()
        p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p2.paragraph_format.space_after = Pt(2)
        r2 = p2.add_run(desc)
        set_default_font(r2, size=9, color=GRAY)


def add_cover_callout(doc):
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    cell = table.rows[0].cells[0]
    set_cell_shading(cell, "EEF4FF")
    p1 = cell.paragraphs[0]
    p1.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r1 = p1.add_run("Proposta de valor")
    set_default_font(r1, size=13, bold=True)
    p1.paragraph_format.space_after = Pt(5)
    for line in [
        "Sistema responsivo para computador e celular",
        "Modulos organizados por setor com foco em uso diario",
        "Mais controle, mais rastreabilidade e menos planilhas soltas",
    ]:
        p = cell.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_after = Pt(3)
        r = p.add_run(line)
        set_default_font(r, size=10, color=GRAY)


def screenshot_slots():
    return [
        ("01_login.png", "Tela de login", "Acesso organizado com proposta visual da plataforma e entrada simplificada para a operacao."),
        ("02_painel_principal.png", "Painel principal", "Visao administrativa por setores, com navegacao clara e distribuicao dos modulos por contexto."),
        ("03_lancamento_diario.png", "Lancamento diario", "Registro da rotina operacional dos equipamentos, horas, status e apontamentos do dia."),
        ("04_abastecimento.png", "Abastecimento", "Controle de diesel e lubrificacao com assinatura, observacoes e leitura de estoque por base."),
        ("05_romaneio_transporte.png", "Romaneio de transporte", "Lancamento da carga com material, quantidade, origem, destino e assinatura para liberar a viagem."),
        ("06_receber_transporte.png", "Recebimento por QR", "Conferencia no destino com leitura do QR do romaneio, agilizando recebimento e rastreabilidade."),
        ("07_saida_materiais.png", "Saidas de materiais", "Retirada de ferramentas, insumos e EPI com assinatura e controle de pendencias."),
        ("08_entrada_materiais.png", "Entrada de materiais", "Entrada central no almoxarifado com estoque atual, registros e filtros para consulta."),
    ]


def mask_box(draw, box, text="", fill="#FFFFFF", outline="#E3EAF6", text_fill=PRIMARY_DARK, radius=16, text_size=18):
    rounded(draw, box, radius=radius, fill=fill, outline=outline, width=2)
    if text:
        text_font = font(text_size, True)
        bbox = draw.textbbox((0, 0), text, font=text_font)
        text_w = bbox[2] - bbox[0]
        text_h = bbox[3] - bbox[1]
        x1, y1, x2, y2 = box
        tx = x1 + ((x2 - x1) - text_w) / 2
        ty = y1 + ((y2 - y1) - text_h) / 2 - 1
        draw.text((tx, ty), text, fill=text_fill, font=text_font)


def sanitize_screenshot(src: Path, dest: Path):
    img = Image.open(src).convert("RGB")
    draw = ImageDraw.Draw(img)
    w, h = img.size
    name = src.name.lower()

    if name == "01_login.png":
        mask_box(draw, (520, 235, 1010, 290), "usuario@empresa.com", fill="#EEF5FF", text_size=20)
        mask_box(draw, (525, 340, 1010, 392), "Senha protegida", fill="#EEF5FF", text_size=20)
    elif name == "02_painel_principal.png":
        mask_box(draw, (1674, 38, 1858, 100), "", fill="#F7FAFF", outline="#F7FAFF", radius=28)
    elif name == "03_lancamento_diario.png":
        mask_box(draw, (630, 18, 750, 58), "", fill="#F3F5F8", outline="#F3F5F8", radius=24)
    elif name == "04_abastecimento.png":
        mask_box(draw, (360, 18, 470, 52), "", fill="#F3F5F8", outline="#F3F5F8", radius=18)
        mask_box(draw, (4, 71, 145, 89), "", fill="#FFFFFF", outline="#FFFFFF", radius=8)
    elif name == "05_romaneio_transporte.png":
        mask_box(draw, (394, 18, 506, 52), "", fill="#F3F5F8", outline="#F3F5F8", radius=18)
        draw.rectangle((10, 228, 205, 244), fill="#FFFFFF")
    elif name == "06_receber_transporte.png":
        mask_box(draw, (1178, 20, 1310, 58), "", fill="#F3F5F8", outline="#F3F5F8", radius=24)
    elif name == "07_saida_materiais.png":
        mask_box(draw, (688, 20, 816, 58), "", fill="#F3F5F8", outline="#F3F5F8", radius=24)
    elif name == "08_entrada_materiais.png":
        mask_box(draw, (688, 20, 816, 58), "", fill="#F3F5F8", outline="#F3F5F8", radius=24)

    dest.parent.mkdir(parents=True, exist_ok=True)
    img.save(dest)


def prepare_screenshots():
    PROCESSED_DIR.mkdir(parents=True, exist_ok=True)
    prepared = []
    for filename, title, desc in screenshot_slots():
        src = SCREENSHOT_DIR / filename
        if not src.exists():
            continue
        dest = PROCESSED_DIR / filename
        sanitize_screenshot(src, dest)
        prepared.append((dest, title, desc))
    return prepared


def create_doc(screenshots):
    doc = Document()
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.left_margin = Inches(0.55)
    section.right_margin = Inches(0.55)
    section.top_margin = Inches(0.55)
    section.bottom_margin = Inches(0.55)

    styles = doc.styles
    styles["Normal"].font.name = "Arial"
    styles["Normal"].font.size = Pt(11)

    header = section.header
    hp = header.paragraphs[0]
    hp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    hr = hp.add_run("Equipamento Gestao | Portfolio do Sistema")
    set_default_font(hr, 9, color=GRAY)

    logo_paragraph = doc.add_paragraph()
    logo_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    logo_run = logo_paragraph.add_run()
    logo_run.add_picture(str(COVER_LOGO_PATH), width=Inches(3.15))
    logo_paragraph.paragraph_format.space_after = Pt(10)
    add_styled_paragraph(doc, "PORTFOLIO COMERCIAL DO SISTEMA", size=11, bold=True, color=GRAY, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=2)
    add_styled_paragraph(doc, "Equipamento Gestao", size=32, bold=True, color=PRIMARY_DARK, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=1)
    add_styled_paragraph(doc, "Controle operacional e administrativo para obras", size=18, bold=True, color=PRIMARY, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=4)
    add_styled_paragraph(
        doc,
        "Uma plataforma pensada para centralizar a rotina da obra, reduzir falhas operacionais e dar visibilidade real ao que acontece no campo e no escritorio.",
        size=11,
        color=GRAY,
        align=WD_ALIGN_PARAGRAPH.CENTER,
        space_after=8,
    )
    add_cover_highlights(doc)
    doc.add_paragraph().paragraph_format.space_after = Pt(6)
    add_cover_callout(doc)
    doc.add_paragraph().paragraph_format.space_after = Pt(8)
    doc.add_picture(str(OVERVIEW_PATH), width=IMAGE_WIDTH)

    add_styled_paragraph(doc, "Resumo executivo", style_name="Heading 1", size=16, bold=True, color=PRIMARY_DARK, space_after=4)
    for bullet in [
        "Sistema web responsivo para operacao diaria no computador e no celular.",
        "Estrutura por setores: Engenharia, Manutencao, Transporte, Almoxarifado, Cadastros e Administrativo Financeiro.",
        "Fluxos com assinatura, QR Code e localizacao para reforcar rastreabilidade.",
        "Permissoes por perfil, base e cidade para reduzir erro e dar governanca.",
    ]:
        add_bullet(doc, bullet)

    doc.add_page_break()
    add_styled_paragraph(doc, "1. Visao geral da plataforma", style_name="Heading 1", size=18, bold=True, color=PRIMARY_DARK, space_after=4)
    add_styled_paragraph(
        doc,
        "O sistema foi desenhado para concentrar toda a rotina da obra em uma unica plataforma: lancamentos do dia, controle tecnico, manutencao, transporte, almoxarifado e gestao administrativa.",
        size=11,
        color=GRAY,
        space_after=10,
    )
    doc.add_picture(str(OVERVIEW_PATH), width=IMAGE_WIDTH)

    table = doc.add_table(rows=1, cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    hdr = table.rows[0].cells
    hdr[0].text = "Area"
    hdr[1].text = "O que entrega"
    set_cell_shading(hdr[0], "EAF2FF")
    set_cell_shading(hdr[1], "EAF2FF")
    for left, right in [
        ("Operacional", "Lancamento diario, RDO, croqui/PDF, transporte, QR, assinatura, manutencao e saidas."),
        ("Administrativo", "Cadastros, controle de equipamentos, funcionarios, bases, obras e configuracoes."),
        ("Administrativo Financeiro", "Plano, historico, acompanhamento financeiro e governanca do cliente."),
    ]:
        row = table.add_row().cells
        row[0].text = left
        row[1].text = right

    doc.add_page_break()
    add_styled_paragraph(doc, "2. Operacao no campo", style_name="Heading 1", size=18, bold=True, color=PRIMARY_DARK, space_after=4)
    add_styled_paragraph(
        doc,
        "No uso operacional, o foco e velocidade. A aplicacao no celular concentra somente os modulos que precisam rodar no dia, reduzindo poluicao visual e acelerando a execucao em campo.",
        size=11,
        color=GRAY,
        space_after=10,
    )
    doc.add_picture(str(OPERATIONAL_PATH), width=IMAGE_WIDTH)
    for bullet in [
        "Lancamento diario de equipamentos e producao.",
        "RDO e producao de campo com croqui/PDF.",
        "Abastecimento, manutencao, saidas de materiais e EPI.",
        "Romaneio de transporte com QR Code, assinatura e geolocalizacao.",
    ]:
        add_bullet(doc, bullet)

    doc.add_page_break()
    add_styled_paragraph(doc, "3. Painel administrativo do cliente", style_name="Heading 1", size=18, bold=True, color=PRIMARY_DARK, space_after=4)
    add_styled_paragraph(
        doc,
        "A camada administrativa organiza os cadastros, relatorios, controles e acompanhamento financeiro do proprio cliente, sem depender de planilhas ou registros paralelos.",
        size=11,
        color=GRAY,
        space_after=10,
    )
    doc.add_picture(str(ADMIN_PATH), width=IMAGE_WIDTH)
    for bullet in [
        "Cadastros estruturais de obras, bases, equipamentos, funcionarios e usuarios operacionais.",
        "Relatorios administrativos para transporte, transferencias, diesel, manutencao e producao.",
        "Financeiro do cliente com visao de plano, vencimento, historico e liberacoes.",
        "Organizacao administrativa para manter operacao, controle e auditoria no mesmo ambiente.",
    ]:
        add_bullet(doc, bullet)

    doc.add_page_break()
    add_styled_paragraph(doc, "4. Mobilidade e experiencia de uso", style_name="Heading 1", size=18, bold=True, color=PRIMARY_DARK, space_after=4)
    add_styled_paragraph(
        doc,
        "A plataforma pode ser instalada no Android como aplicativo, com navegacao simplificada, icone proprio e experiencia focada em operacao rapida.",
        size=11,
        color=GRAY,
        space_after=10,
    )
    doc.add_picture(str(MOBILE_PATH), width=IMAGE_WIDTH)
    for bullet in [
        "Instalacao como app pela tela inicial do celular.",
        "Logo personalizada, experiencia PWA e tela inicial mais profissional.",
        "Setores visiveis sem arrastar para o lado e modulos organizados por contexto.",
    ]:
        add_bullet(doc, bullet)

    doc.add_page_break()
    add_styled_paragraph(doc, "5. Diferenciais competitivos", style_name="Heading 1", size=18, bold=True, color=PRIMARY_DARK, space_after=6)
    for bullet in [
        "Sistema pensado para a rotina real de obra, nao apenas para escritorio.",
        "Rastreabilidade operacional com QR Code, assinatura, horario e localizacao.",
        "Separacao por perfil, base e cidade para reduzir erro e dar governanca.",
        "Estrutura pronta para seguir evoluindo com novos modulos sem desmontar o sistema.",
    ]:
        add_bullet(doc, bullet)

    if screenshots:
        doc.add_page_break()
        add_styled_paragraph(doc, "6. Galeria de telas reais", style_name="Heading 1", size=18, bold=True, color=PRIMARY_DARK, space_after=4)
        add_styled_paragraph(
            doc,
            "A seguir, telas reais do sistema em uso, destacando operacao, transporte, almoxarifado e controles diarios.",
            size=11,
            color=GRAY,
            space_after=10,
        )
        for idx, (img_path, title, desc) in enumerate(screenshots):
            add_styled_paragraph(doc, title, style_name="Heading 2", size=14, bold=True, color=PRIMARY_DARK, space_after=2)
            add_styled_paragraph(doc, desc, size=10, color=GRAY, space_after=6)
            doc.add_picture(str(img_path), width=IMAGE_WIDTH)
            if idx != len(screenshots) - 1:
                if idx % 2 == 1:
                    doc.add_page_break()
                else:
                    doc.add_paragraph().paragraph_format.space_after = Pt(10)

    footer = section.footer
    fp = footer.paragraphs[0]
    fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    fr = fp.add_run("Equipamento Gestao | Controle operacional e administrativo")
    set_default_font(fr, 9, color=GRAY)

    doc.save(DOCX_PATH)


def main():
    OUT_DIR.mkdir(parents=True, exist_ok=True)
    PROCESSED_DIR.mkdir(parents=True, exist_ok=True)
    screenshots = prepare_screenshots()
    create_doc(screenshots)
    print(DOCX_PATH)


if __name__ == "__main__":
    main()
